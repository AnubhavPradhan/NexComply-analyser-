"""
Policy processor module.

Handles loading and processing of policy documents in various formats.
"""
from typing import Optional, List
from pathlib import Path
import PyPDF2
import docx
from loguru import logger


class PolicyProcessor:
    """
    Process and extract text from policy documents.
    """
    
    def __init__(self):
        """Initialize the policy processor."""
        pass
    
    def load_document(self, file_path: str) -> str:
        """
        Load and extract text from a document.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            str: Extracted text content.
        """
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return ""
        
        extension = path.suffix.lower()
        
        if extension == ".pdf":
            return self._extract_from_pdf(path)
        elif extension in [".docx", ".doc"]:
            return self._extract_from_docx(path)
        elif extension == ".txt":
            return self._extract_from_txt(path)
        else:
            logger.warning(f"Unsupported file format: {extension}")
            return ""
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file.
            
        Returns:
            str: Extracted text.
        """
        try:
            logger.info(f"Extracting text from PDF: {file_path}")
            text = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            logger.info(f"Extracted {len(text)} characters from PDF")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
            return ""
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file.
            
        Returns:
            str: Extracted text.
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")
            doc = docx.Document(file_path)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
            return ""
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file.
            
        Returns:
            str: Extracted text.
        """
        try:
            logger.info(f"Reading text from TXT: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Read {len(text)} characters from TXT")
            return text
            
        except Exception as e:
            logger.error(f"Error reading TXT file: {e}")
            return ""
    
    def load_from_uploaded_file(self, uploaded_file) -> str:
        """
        Load document from Streamlit uploaded file.
        
        Args:
            uploaded_file: Streamlit UploadedFile object.
            
        Returns:
            str: Extracted text content.
        """
        try:
            file_extension = Path(uploaded_file.name).suffix.lower()
            
            if file_extension == ".pdf":
                return self._extract_from_pdf_bytes(uploaded_file.read())
            elif file_extension in [".docx", ".doc"]:
                return self._extract_from_docx_bytes(uploaded_file.read())
            elif file_extension == ".txt":
                return uploaded_file.read().decode('utf-8')
            else:
                logger.warning(f"Unsupported file format: {file_extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Error loading uploaded file: {e}")
            return ""
    
    def _extract_from_pdf_bytes(self, pdf_bytes: bytes) -> str:
        """
        Extract text from PDF bytes.
        
        Args:
            pdf_bytes: PDF file content as bytes.
            
        Returns:
            str: Extracted text.
        """
        try:
            from io import BytesIO
            pdf_file = BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting from PDF bytes: {e}")
            return ""
    
    def _extract_from_docx_bytes(self, docx_bytes: bytes) -> str:
        """
        Extract text from DOCX bytes.
        
        Args:
            docx_bytes: DOCX file content as bytes.
            
        Returns:
            str: Extracted text.
        """
        try:
            from io import BytesIO
            docx_file = BytesIO(docx_bytes)
            doc = docx.Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX bytes: {e}")
            return ""
